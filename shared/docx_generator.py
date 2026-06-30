"""
DOCX generation utility for Finance documents.
Returns a python-docx Document object — caller streams it as HTTP response.
"""
import os
from pathlib import Path
from decimal import Decimal
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime

# Single source of truth: the template lives in the NAS shared folder so it can be
# edited directly in Word (N: share). Falls back to code-built output if unavailable.
TEMPLATE_DIR = Path('/mnt/data/1os/shared/doc_templates')


def _get_tenant():
    from django.apps import apps
    Tenant = apps.get_model('accounts', 'Tenant')
    return Tenant.objects.first()


def _fmt_date(d):
    return d.strftime('%d %B %Y') if d else ''


def _fmt_money(v):
    return f"{(v or Decimal('0')):,.2f}"


def _fmt_qty(v):
    v = v or Decimal('0')
    return str(int(v)) if v == v.to_integral_value() else f"{v:,.2f}"


def _media_path(field):
    """Resolve a tenant image field (FileField or plain path string) to an existing abs path."""
    from django.conf import settings
    if not field:
        return None
    try:
        p = field.path                       # ImageField / FileField
    except (ValueError, AttributeError):
        p = os.path.join(settings.MEDIA_ROOT, str(field))  # CharField path
    return p if os.path.exists(p) else None


def _header(doc, title, doc_no, issue_date, client_name, client_address=None, extra_rows=None, tenant=None):
    """Add company header + document title + client block."""
    t = tenant or _get_tenant()
    company_name    = t.name if t else 'Simply Engineering'
    company_address = t.address if t else ''
    company_tel     = t.phone if t else ''
    uen_label = f'UEN: {t.uen}' if t and t.uen else ''
    gst_label = f'GST Reg No: {t.gst_number}' if t and t.gst_registered and t.gst_number else ''

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.bold = True
    run.font.size = Pt(14)

    for line in filter(None, [company_address, company_tel, uen_label, gst_label]):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Document title
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)

    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'

    def add_row(label1, val1, label2='', val2=''):
        row = table.add_row().cells
        row[0].text = label1
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(val1)
        row[2].text = label2
        if label2:
            row[2].paragraphs[0].runs[0].bold = True
        row[3].text = str(val2)

    add_row(f'{title} No.', doc_no, 'Date', str(issue_date))
    add_row('Client', client_name)
    if client_address:
        add_row('Address', client_address)
    if extra_rows:
        for r in extra_rows:
            add_row(*r)

    doc.add_paragraph()


def _items_table(doc, headers, rows, totals=None):
    """Add a line items table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val) if val is not None else ''

    # Totals
    if totals:
        for label, value in totals:
            row = table.add_row().cells
            span = len(headers) - 2
            row[span].text = label
            row[span].paragraphs[0].runs[0].bold = True
            row[span].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row[span + 1].text = f'${value:,.2f}'


def generate_quotation(quotation):
    """Render the quotation from the Word template if present, else fall back to code-built."""
    template = TEMPLATE_DIR / 'quotation.docx'
    if not template.exists():
        return _generate_quotation_codebuilt(quotation)
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return _generate_quotation_codebuilt(quotation)

    import copy
    import jinja2
    from docx.oxml.ns import qn

    tenant = _get_tenant()
    tpl = DocxTemplate(str(template))
    # autoescape so field values containing & < > render correctly (docxtpl drops them otherwise)
    jenv = jinja2.Environment(autoescape=True)

    # swap embedded images per-tenant: image1 = logo, image2 = signature
    if tenant:
        logo = _media_path(getattr(tenant, 'logo', None))
        if logo:
            try:
                tpl.replace_pic('image1.jpeg', logo)
            except Exception:
                pass
        sig = _media_path(getattr(tenant, 'signatory_file', None))
        if sig:
            try:
                tpl.replace_pic('image2.png', sig)
            except Exception:
                pass

    signatory = (tenant.signatory_name if tenant and tenant.signatory_name else '') \
        or (quotation.prepared_by.full_name if quotation.prepared_by else '')

    tpl.render({
        'title':            quotation.title or '',
        'quote_no':         quotation.quote_no,
        'issue_date':       _fmt_date(quotation.issue_date),
        'valid_until':      _fmt_date(quotation.valid_until),
        'status':           quotation.get_status_display(),
        'customer_name':    quotation.client_name or '',
        'customer_contact': quotation.client_contact or '',
        'customer_address': quotation.client_address or '',
        'company_address':  (tenant.address if tenant else '') or '',
        'company_phone':    (tenant.phone if tenant else '') or '',
        'company_email':    (tenant.email if tenant else '') or '',
        'company_website':  (tenant.site_url if tenant else '') or '',
        'company_uen':      (tenant.uen if tenant else '') or '',
        'company_gst_no':   (tenant.gst_number if tenant and tenant.gst_registered else '') or '',
        'subtotal':         _fmt_money(quotation.subtotal),
        'gst_amount':       _fmt_money(quotation.gst_amount),
        'total':            _fmt_money(quotation.total),
        'notes':            quotation.notes or '',
        'signatory_name':         signatory,
        'signatory_designation':  (tenant.signatory_designation if tenant else '') or '',
    }, jinja_env=jenv)

    # fill the repeating items row from sentinel markers (docxtpl {%tr%} is unreliable here)
    doc = tpl.docx
    SENT = {'#NO#': 'no', '#DESC#': 'description', '#QTY#': 'qty',
            '#UNIT#': 'unit', '#PRICE#': 'unit_price', '#AMOUNT#': 'amount'}
    rows = [{
        'no': i, 'description': it.description, 'qty': _fmt_qty(it.qty),
        'unit': it.unit or '', 'unit_price': _fmt_money(it.unit_price),
        'amount': _fmt_money(it.amount),
    } for i, it in enumerate(quotation.items.all(), 1)]

    for table in doc.tables:
        marker = next((r for r in table.rows if any('#NO#' in c.text for c in r.cells)), None)
        if marker is None:
            continue
        tr = marker._tr
        parent = tr.getparent()
        idx = list(parent).index(tr)
        for item in rows:
            new = copy.deepcopy(tr)
            for t in new.iter(qn('w:t')):
                if t.text:
                    for s, k in SENT.items():
                        if s in t.text:
                            t.text = t.text.replace(s, str(item[k]))
            parent.insert(idx, new)
            idx += 1
        parent.remove(tr)
        break
    return doc


def _generate_quotation_codebuilt(quotation):
    doc = Document()
    tenant = _get_tenant()
    _header(
        doc, 'QUOTATION',
        quotation.quote_no,
        quotation.issue_date,
        quotation.client_name,
        quotation.client_address,
        extra_rows=[
            ('Valid Until', str(quotation.valid_until or '—'), 'Status', quotation.status.upper()),
        ],
        tenant=tenant,
    )

    headers = ['No.', 'Description', 'Unit', 'Qty', 'Unit Price (S$)', 'Amount (S$)']
    rows = []
    for i, item in enumerate(quotation.items.all(), 1):
        rows.append([i, item.description, item.unit or '', item.qty, f'{item.unit_price:,.2f}', f'{item.amount:,.2f}'])

    totals = [
        ('Subtotal', quotation.subtotal),
        ('GST (9%)', quotation.gst_amount),
        ('Total', quotation.total),
    ]
    _items_table(doc, headers, rows, totals)

    if quotation.notes:
        doc.add_paragraph()
        doc.add_paragraph(f'Notes: {quotation.notes}')

    _footer(doc, quotation.prepared_by, tenant=tenant)
    return doc


def generate_invoice(invoice):
    doc = Document()
    tenant = _get_tenant()
    _header(
        doc, 'INVOICE',
        invoice.invoice_no,
        invoice.issue_date,
        invoice.client_name,
        extra_rows=[
            ('Due Date', str(invoice.due_date), 'Status', invoice.status.upper()),
        ],
        tenant=tenant,
    )

    headers = ['No.', 'Description', 'Unit', 'Qty', 'Unit Price (S$)', 'Amount (S$)']
    rows = []
    for i, item in enumerate(invoice.items.all(), 1):
        rows.append([i, item.description, item.unit or '', item.qty, f'{item.unit_price:,.2f}', f'{item.amount:,.2f}'])

    totals = [
        ('Subtotal', invoice.subtotal),
        ('GST (9%)', invoice.gst_amount),
        ('Total', invoice.total),
        ('Paid', invoice.paid_amount),
        ('Balance Due', invoice.balance_due),
    ]
    _items_table(doc, headers, rows, totals)

    if invoice.notes:
        doc.add_paragraph()
        doc.add_paragraph(f'Notes: {invoice.notes}')

    _footer(doc, tenant=tenant)
    return doc


def generate_delivery_order(do):
    doc = Document()
    tenant = _get_tenant()
    _header(
        doc, 'DELIVERY ORDER',
        do.do_no,
        do.issue_date,
        do.client_name,
        do.delivery_address,
        extra_rows=[
            ('Delivery Date', str(do.delivery_date or '—'), 'Status', do.status.upper()),
            ('Delivered By', do.delivered_by or '—', 'Received By', do.received_by or '—'),
        ],
        tenant=tenant,
    )

    headers = ['No.', 'Description', 'Unit', 'Qty', 'Remarks']
    rows = []
    for i, item in enumerate(do.items.all(), 1):
        rows.append([i, item.description, item.unit or '', item.qty, item.remarks or ''])

    _items_table(doc, headers, rows)

    doc.add_paragraph()
    # Signature block
    sig = doc.add_table(rows=3, cols=2)
    sig.rows[0].cells[0].text = 'Delivered by (Signature / Name):'
    sig.rows[0].cells[1].text = 'Received by (Signature / Name / Date):'
    sig.rows[2].cells[0].text = f'{do.delivered_by or ""}'
    sig.rows[2].cells[1].text = f'{do.received_by or ""}'

    if do.notes:
        doc.add_paragraph()
        doc.add_paragraph(f'Notes: {do.notes}')

    _footer(doc, do.prepared_by, tenant=tenant)
    return doc


def _footer(doc, prepared_by=None, tenant=None):
    t = tenant or _get_tenant()
    signatory = (t.signatory_name if t and t.signatory_name else '') or (prepared_by.full_name if prepared_by else '')
    designation = t.signatory_designation if t and t.signatory_designation else ''

    doc.add_paragraph()

    if signatory or designation:
        p = doc.add_paragraph()
        p.add_run('Authorised by: ')
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)
        p.add_run(f'{signatory}  {designation}'.strip())
        p.runs[-1].font.size = Pt(9)

        if t and t.signatory_file:
            try:
                from docx.shared import Inches
                doc.add_picture(t.signatory_file.path, width=Inches(1.5))
            except Exception:
                pass

    p = doc.add_paragraph()
    p.add_run(f'Generated by 1OS on {datetime.date.today()}')
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
